#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para converter Markdown para DOCX
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def converter_markdown_para_docx(arquivo_md, arquivo_docx):
    """Converte arquivo Markdown para DOCX"""
    
    # Ler arquivo Markdown
    with open(arquivo_md, 'r', encoding='utf-8') as f:
        conteudo = f.read()
    
    # Criar documento DOCX
    doc = Document()
    
    # Processar linhas
    linhas = conteudo.split('\n')
    em_codigo = False
    codigo_linhas = []
    
    for linha in linhas:
        linha_stripped = linha.strip()
        
        # Bloco de código
        if linha_stripped.startswith('```'):
            if em_codigo:
                # Fim do bloco de código
                paragrafo = doc.add_paragraph()
                paragrafo.style = 'No Spacing'
                for codigo_linha in codigo_linhas:
                    run = paragrafo.add_run(codigo_linha + '\n')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                codigo_linhas = []
                em_codigo = False
            else:
                # Início do bloco de código
                em_codigo = True
            continue
        
        # Se está em bloco de código
        if em_codigo:
            codigo_linhas.append(linha_stripped)
            continue
        
        # Títulos (H1, H2, H3, etc.)
        if linha_stripped.startswith('# '):
            paragrafo = doc.add_heading(linha_stripped[2:], level=1)
            continue
        elif linha_stripped.startswith('## '):
            paragrafo = doc.add_heading(linha_stripped[3:], level=2)
            continue
        elif linha_stripped.startswith('### '):
            paragrafo = doc.add_heading(linha_stripped[4:], level=3)
            continue
        
        # Linhas vazias
        if not linha_stripped:
            doc.add_paragraph()
            continue
        
        # Listas com ✅
        if linha_stripped.startswith('✅'):
            paragrafo = doc.add_paragraph(linha_stripped, style='List Bullet')
            continue
        
        # Listas com hífen
        if linha_stripped.startswith('- '):
            paragrafo = doc.add_paragraph(linha_stripped[2:], style='List Bullet')
            continue
        
        # Linhas horizontais
        if linha_stripped.startswith('---'):
            doc.add_paragraph('_' * 50)
            continue
        
        # Texto normal
        paragrafo = doc.add_paragraph(linha_stripped)
    
    # Salvar documento
    doc.save(arquivo_docx)
    print(f"✅ Documento criado com sucesso: {arquivo_docx}")

if __name__ == '__main__':
    arquivo_md = r'c:\Users\Julia Zanette\Desktop\M3 - Sistemas Distribuidos\ATIVIDADE_RMI_SOLUCAO.md'
    arquivo_docx = r'c:\Users\Julia Zanette\Desktop\M3 - Sistemas Distribuidos\ATIVIDADE_RMI_SOLUCAO.docx'
    
    try:
        converter_markdown_para_docx(arquivo_md, arquivo_docx)
    except Exception as e:
        print(f"❌ Erro: {e}")
